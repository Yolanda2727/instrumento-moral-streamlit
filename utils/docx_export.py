from __future__ import annotations

from io import BytesIO
from typing import Any, Dict, Iterable, Sequence, Tuple

import pandas as pd
from docx import Document


def _safe_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and pd.isna(value):
        return ""
    return str(value)


def _add_key_value_table(document: Document, rows: Sequence[Tuple[str, Any]], *, title: str | None = None) -> None:
    if title:
        document.add_heading(title, level=2)
    table = document.add_table(rows=max(len(rows), 1), cols=2)
    table.style = "Table Grid"
    if not rows:
        table.cell(0, 0).text = "mensaje"
        table.cell(0, 1).text = "Sin datos disponibles"
        return
    for idx, (key, value) in enumerate(rows):
        table.cell(idx, 0).text = _safe_text(key)
        table.cell(idx, 1).text = _safe_text(value)


def _add_dataframe_table(
    document: Document,
    df: pd.DataFrame,
    *,
    title: str,
    max_rows: int = 500,
) -> None:
    document.add_heading(title, level=2)
    if df is None or df.empty:
        document.add_paragraph("Sin datos disponibles.")
        return

    export_df = df.head(max_rows).copy()
    table = document.add_table(rows=len(export_df) + 1, cols=len(export_df.columns))
    table.style = "Table Grid"

    for j, col in enumerate(export_df.columns):
        table.cell(0, j).text = _safe_text(col)

    for i in range(len(export_df)):
        for j, col in enumerate(export_df.columns):
            table.cell(i + 1, j).text = _safe_text(export_df.iloc[i][col])

    if len(df) > max_rows:
        document.add_paragraph(f"(Mostrando {max_rows} filas de {len(df)}.)")


def _add_bullet_list(document: Document, items: Iterable[Any]) -> None:
    for item in items:
        text = _safe_text(item).strip()
        if not text:
            continue
        document.add_paragraph(text, style="List Bullet")


def build_individual_report_docx_bytes(
    *,
    app_title: str,
    app_brand_line: str,
    author_name: str,
    author_credentials: Sequence[str],
    main_function: str,
    generated_at: str,
    participant_rows: Sequence[Tuple[str, Any]],
    metric_rows: Sequence[Tuple[str, Any]],
    narrative_summary: str,
    recommendations: Sequence[str],
    framework_scores_df: pd.DataFrame,
    stage_scores_df: pd.DataFrame,
    choice_detail_df: pd.DataFrame,
    interpretation_result: Dict[str, Any] | None = None,
    interpretation_note: str | None = None,
) -> bytes:
    document = Document()

    document.add_heading(f"{app_title} — Reporte individual", level=0)
    document.add_paragraph(app_brand_line)
    document.add_paragraph(f"Autor: {author_name}")
    if author_credentials:
        document.add_paragraph("Credenciales: " + "; ".join([_safe_text(x) for x in author_credentials if _safe_text(x)]))
    document.add_paragraph(f"Función principal: {main_function}")
    document.add_paragraph(f"Generado en: {generated_at}")

    _add_key_value_table(document, participant_rows, title="Datos del participante")
    _add_key_value_table(document, metric_rows, title="Métricas principales")

    document.add_heading("Síntesis interpretativa", level=2)
    document.add_paragraph(_safe_text(narrative_summary))

    document.add_heading("Recomendaciones de mejora argumentativa", level=2)
    _add_bullet_list(document, recommendations)

    _add_dataframe_table(document, framework_scores_df, title="Tabla de marcos éticos")
    _add_dataframe_table(document, stage_scores_df, title="Tabla de estadios morales")
    _add_dataframe_table(document, choice_detail_df, title="Detalle de respuestas")

    document.add_heading("Análisis IA integrado", level=2)
    if interpretation_result:
        for key in ["resumen_ejecutivo", "interpretacion_etica", "interpretacion_legal", "interpretacion_bioetica"]:
            value = interpretation_result.get(key)
            if _safe_text(value).strip():
                document.add_heading(key.replace("_", " ").title(), level=3)
                document.add_paragraph(_safe_text(value))

        for list_key, title in [
            ("consideraciones_clave", "Consideraciones clave"),
            ("fortalezas_argumentativas", "Fortalezas argumentativas"),
            ("debilidades_argumentativas", "Debilidades argumentativas"),
            ("recomendaciones_formativas", "Recomendaciones formativas"),
        ]:
            items = interpretation_result.get(list_key) or []
            if isinstance(items, list) and items:
                document.add_heading(title, level=3)
                _add_bullet_list(document, items)

        riesgos = interpretation_result.get("riesgos") or []
        if isinstance(riesgos, list) and riesgos:
            document.add_heading("Riesgos", level=3)
            riesgos_df = pd.DataFrame(riesgos)
            _add_dataframe_table(document, riesgos_df, title="Riesgos (tabla)", max_rows=200)

        tabla = interpretation_result.get("tabla_analitica") or []
        if isinstance(tabla, list) and tabla:
            document.add_heading("Tabla analítica", level=3)
            tabla_df = pd.DataFrame(tabla)
            _add_dataframe_table(document, tabla_df, title="Tabla analítica (tabla)", max_rows=200)
    elif interpretation_note:
        document.add_paragraph(_safe_text(interpretation_note))
    else:
        document.add_paragraph("No hay interpretación IA disponible.")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
